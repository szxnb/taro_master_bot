from docx import Document
from docx.shared import Pt, RGBColor
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import card_detail
import send_email
from docx2pdf import convert

# str 问题, dic 三个卡牌的字典, dic 四个回答的字典 
def generate(qusetion, card1, card2, card3, answers, recipient_email):

    user_question = qusetion
    master_words =  '''Before we begin the interpretation, I would like to explain some important points to you. Tarot cards are not tools for predicting the future, but rather tools for guiding us to delve deeper into various situations in life. Therefore, when receiving the reading, please maintain an open mindset and trust your intuition. Our goal is to provide you with inspiration and guidance through the guidance of the tarot cards, helping you better deal with challenges and decisions in life. Alright, you can start reading now!
    '''
    # Name
    card_first_name = 'CardⅠ ' + card_detail.card_detail[card1][0]
    card_second_name = 'CardⅡ ' + card_detail.card_detail[card2][0]
    card_third_name = 'CardⅢ ' + card_detail.card_detail[card3][0]
    # Image
    card_first_image = '../assets/' + card_detail.card_detail[card1][1]
    card_second_image = '../assets/' + card_detail.card_detail[card2][1]
    card_third_image = '../assets/' + card_detail.card_detail[card3][1]
    # Description
    card_first_description = card_detail.card_detail[card1][2]
    card_second_description = card_detail.card_detail[card2][2]
    card_third_description = card_detail.card_detail[card3][2]

    summarize = answers['result_overview']
    love_suggestion = answers['result_love']
    career_suggestion = answers['result_career']
    wealth_suggestion = answers['result_finances']

    # 读取文档
    doc = Document('../templates/template.docx')

    def limit_run_length(run, max_length):
        if len(run.text) > max_length:
            run.text = run.text[:max_length - 3] + '...'

    # 遍历文档中的所有表格
    for table in doc.tables:
        # 遍历表格中的所有行
        for row in table.rows:
            # 遍历行中的所有单元格
            for cell in row.cells:
                # 遍历单元格中的所有段落
                for paragraph in cell.paragraphs:
                    if '{{#date}}' in paragraph.text:
                        # 设置段落的对齐方式
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # 获取当前时间
                        current_time = datetime.now()
                        # 将时间格式化为 "Month day, year" 格式
                        formatted_time = current_time.strftime("%B %d, %Y")
                        # 遍历段落中的所有运行
                        for run in paragraph.runs:
                            # 替换运行中的文字
                            run.text = run.text.replace('{{#date}}', formatted_time)
                            # 限制运行的长度
                            limit_run_length(run, 300)
                    if '{{#user_question}}' in paragraph.text:
                        # 设置段落的对齐方式
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # 设置段落的行间距
                        # 遍历段落中的所有运行
                        for run in paragraph.runs:
                            # 替换运行中的文字
                            run.text = run.text.replace('{{#user_question}}', user_question)
                            # 限制运行的长度
                            limit_run_length(run, 300)
                    if '{{#master_words}}' in paragraph.text:
                        # 设置段落的对齐方式
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        # 遍历段落中的所有运行
                        for run in paragraph.runs:
                            # 替换运行中的文字
                            run.text = run.text.replace('{{#master_words}}', master_words)
                            # 限制运行的长度
                            limit_run_length(run, 600)
                    if '{{#card_first_name}}' in paragraph.text:
                            # 设置段落的对齐方式
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            original_style = paragraph.style
                            original_runs = paragraph.runs
                            paragraph.text = paragraph.text.replace('{{#card_first_name}}', card_first_name)
                            paragraph.style = original_style
                            for run in paragraph.runs:
                                run.bold = original_runs[0].bold
                                run.italic = original_runs[0].italic
                                run.underline = original_runs[0].underline
                                run.font.size = original_runs[0].font.size
                                run.font.name = original_runs[0].font.name
                                run.font.color.rgb = original_runs[0].font.color.rgb
                    if '{{#card_first_description}}' in paragraph.text:
                            # 设置段落的对齐方式
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            original_style = paragraph.style
                            original_runs = paragraph.runs
                            paragraph.text = paragraph.text.replace('{{#card_first_description}}', card_first_description)
                            paragraph.style = original_style
                            for run in paragraph.runs:
                                run.bold = original_runs[0].bold
                                run.italic = original_runs[0].italic
                                run.underline = original_runs[0].underline
                                run.font.size = original_runs[0].font.size
                                run.font.name = original_runs[0].font.name
                                run.font.color.rgb = original_runs[0].font.color.rgb
                    if '{{#card_second_name}}' in paragraph.text:
                            # 设置段落的对齐方式
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            original_style = paragraph.style
                            original_runs = paragraph.runs
                            paragraph.text = paragraph.text.replace('{{#card_second_name}}', card_second_name)
                            paragraph.style = original_style
                            for run in paragraph.runs:
                                run.bold = original_runs[0].bold
                                run.italic = original_runs[0].italic
                                run.underline = original_runs[0].underline
                                run.font.size = original_runs[0].font.size
                                run.font.name = original_runs[0].font.name
                                run.font.color.rgb = original_runs[0].font.color.rgb
                    if '{{#card_second_description}}' in paragraph.text:
                            # 设置段落的对齐方式
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            original_style = paragraph.style
                            original_runs = paragraph.runs
                            paragraph.text = paragraph.text.replace('{{#card_second_description}}', card_second_description)
                            paragraph.style = original_style
                            for run in paragraph.runs:
                                run.bold = original_runs[0].bold
                                run.italic = original_runs[0].italic
                                run.underline = original_runs[0].underline
                                run.font.size = original_runs[0].font.size
                                run.font.name = original_runs[0].font.name
                                run.font.color.rgb = original_runs[0].font.color.rgb
                    if '{{#card_third_name}}' in paragraph.text:
                            # 设置段落的对齐方式
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            original_style = paragraph.style
                            original_runs = paragraph.runs
                            paragraph.text = paragraph.text.replace('{{#card_third_name}}', card_third_name)
                            paragraph.style = original_style
                            for run in paragraph.runs:
                                run.bold = original_runs[0].bold
                                run.italic = original_runs[0].italic
                                run.underline = original_runs[0].underline
                                run.font.size = original_runs[0].font.size
                                run.font.name = original_runs[0].font.name
                                run.font.color.rgb = original_runs[0].font.color.rgb
                    if '{{#card_third_description}}' in paragraph.text:
                            # 设置段落的对齐方式
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            original_style = paragraph.style
                            original_runs = paragraph.runs
                            paragraph.text = paragraph.text.replace('{{#card_third_description}}', card_third_description)
                            paragraph.style = original_style
                            for run in paragraph.runs:
                                run.bold = original_runs[0].bold
                                run.italic = original_runs[0].italic
                                run.underline = original_runs[0].underline
                                run.font.size = original_runs[0].font.size
                                run.font.name = original_runs[0].font.name
                                run.font.color.rgb = original_runs[0].font.color.rgb
                    if '{{#summarize}}' in paragraph.text:
                            # 设置段落的对齐方式
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            original_style = paragraph.style
                            original_runs = paragraph.runs
                            paragraph.text = paragraph.text.replace('{{#summarize}}', summarize)
                            # 限制运行的长度
                            limit_run_length(run, 5000)
                            paragraph.style = original_style
                            for run in paragraph.runs:
                                run.bold = original_runs[0].bold
                                run.italic = original_runs[0].italic
                                run.underline = original_runs[0].underline
                                run.font.size = original_runs[0].font.size
                                run.font.name = original_runs[0].font.name
                                run.font.color.rgb = original_runs[0].font.color.rgb
                    if '{{#love_suggestion}}' in paragraph.text:
                            # 设置段落的对齐方式
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            original_style = paragraph.style
                            original_runs = paragraph.runs
                            paragraph.text = paragraph.text.replace('{{#love_suggestion}}', love_suggestion)
                            # 限制运行的长度
                            limit_run_length(run, 5000)
                            paragraph.style = original_style
                            for run in paragraph.runs:
                                run.bold = original_runs[0].bold
                                run.italic = original_runs[0].italic
                                run.underline = original_runs[0].underline
                                run.font.size = original_runs[0].font.size
                                run.font.name = original_runs[0].font.name
                                run.font.color.rgb = original_runs[0].font.color.rgb
                    if '{{#career_suggestion}}' in paragraph.text:
                            # 设置段落的对齐方式
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            original_style = paragraph.style
                            original_runs = paragraph.runs
                            paragraph.text = paragraph.text.replace('{{#career_suggestion}}', career_suggestion)
                            # 限制运行的长度
                            limit_run_length(run, 5000)
                            paragraph.style = original_style
                            for run in paragraph.runs:
                                run.bold = original_runs[0].bold
                                run.italic = original_runs[0].italic
                                run.underline = original_runs[0].underline
                                run.font.size = original_runs[0].font.size
                                run.font.name = original_runs[0].font.name
                                run.font.color.rgb = original_runs[0].font.color.rgb
                    if '{{#wealth_suggestion}}' in paragraph.text:
                            # 设置段落的对齐方式
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            original_style = paragraph.style
                            original_runs = paragraph.runs
                            paragraph.text = paragraph.text.replace('{{#wealth_suggestion}}', wealth_suggestion)
                            # 限制运行的长度
                            limit_run_length(run, 5000)
                            paragraph.style = original_style
                            for run in paragraph.runs:
                                run.bold = original_runs[0].bold
                                run.italic = original_runs[0].italic
                                run.underline = original_runs[0].underline
                                run.font.size = original_runs[0].font.size
                                run.font.name = original_runs[0].font.name
                                run.font.color.rgb = original_runs[0].font.color.rgb
                    if '{{#card_first_image}}' in paragraph.text:
                            # 设置段落的对齐方式
                            paragraph.style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            original_style = paragraph.style
                            original_runs = paragraph.runs
                            # 清空单元格中的文本
                            cell.text = ''
                            # 向单元格中插入图片
                            cell.paragraphs[0].add_run().add_picture(card_first_image, width = Inches(2), height = Inches(3))
                            # paragraph.style = original_style
                            for run in paragraph.runs:
                                run.bold = original_runs[0].bold
                                run.italic = original_runs[0].italic
                                run.underline = original_runs[0].underline
                                run.font.size = original_runs[0].font.size
                                run.font.name = original_runs[0].font.name
                                run.font.color.rgb = original_runs[0].font.color.rgb
                    if '{{#card_second_image}}' in paragraph.text:
                            # 设置段落的对齐方式
                            paragraph.style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            original_style = paragraph.style
                            original_runs = paragraph.runs
                            # 清空单元格中的文本
                            cell.text = ''
                            # 向单元格中插入图片
                            cell.paragraphs[0].add_run().add_picture(card_second_image, width = Inches(2), height = Inches(3))
                            # paragraph.style = original_style
                            for run in paragraph.runs:
                                run.bold = original_runs[0].bold
                                run.italic = original_runs[0].italic
                                run.underline = original_runs[0].underline
                                run.font.size = original_runs[0].font.size
                                run.font.name = original_runs[0].font.name
                                run.font.color.rgb = original_runs[0].font.color.rgb
                    if '{{#card_third_image}}' in paragraph.text:
                            # 设置段落的对齐方式
                            paragraph.style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            original_style = paragraph.style
                            original_runs = paragraph.runs
                            # 清空单元格中的文本
                            cell.text = ''
                            # 向单元格中插入图片
                            cell.paragraphs[0].add_run().add_picture(card_third_image, width = Inches(2), height = Inches(3))
                            # paragraph.style = original_style
                            for run in paragraph.runs:
                                run.bold = original_runs[0].bold
                                run.italic = original_runs[0].italic
                                run.underline = original_runs[0].underline
                                run.font.size = original_runs[0].font.size
                                run.font.name = original_runs[0].font.name
                                run.font.color.rgb = original_runs[0].font.color.rgb
                    
    # 格式化当前时间
    formatted_time = datetime.now().strftime("%Y%m%d%H%M%S")
    report_name = "reports/TarotReport_" + recipient_email + "_" + formatted_time
    print(report_name + "生成完成!")
    # 保存修改后的文档
    doc.save("../" + report_name + '.docx')
    input_docx = "../" + report_name + ".docx"
    output_pdf = "../" + report_name + ".pdf"
    convert(input_docx, output_pdf)

    file_locate = output_pdf
    send_email.send_email(file_locate, recipient_email)

# answers = {
#             'result_overview' : 'result_overview',
#             'result_love' : 'result_love',
#             'result_career' : 'result_career',
#             'result_finances' : 'result_finances'
#         }
# generate_docx('can i marry this year?','Star', "Moon", "Sun", answers)













